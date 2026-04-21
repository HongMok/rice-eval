#!/usr/bin/env python3
"""
Parse C-PEP-3 test manual DOCX and generate cpep-config-v2.csv.

Format rules:
- 评分从高到低: a=2(P/A), b=1(E/M), c=0(F/S)
- 一级分类统一"发展"
- 双维度项目(发展+病理学)拆成2行，共享排序号/项目名/操作描述/材料
- 二级分类用原始子领域名
- All fields QUOTE_ALL, UTF-8
"""

import csv
import re
from docx import Document

DOCX_FILE = "data/C-PEP/C-PEP-3测试手册&得分结果图.docx"
OUTPUT = "data/cpep-config-v2.csv"

V2_HEADER = [
    "排序", "一级分类", "二级分类", "评估项目", "操作描述", "所需材料",
    "适用年龄", "是否必答",
    "评分_a_分值", "评分_a_标签", "评分_a_说明",
    "评分_b_分值", "评分_b_标签", "评分_b_说明",
    "评分_c_分值", "评分_c_标签", "评分_c_说明",
    "评分_d_分值", "评分_d_标签", "评分_d_说明",
    "评分_e_分值", "评分_e_标签", "评分_e_说明",
]


def clean_text(text):
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def clean_chinese(text):
    text = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff(（])', '', text)
    text = re.sub(r'(?<=[(（])\s+', '', text)
    text = re.sub(r'\s+(?=[)）])', '', text)
    return text


def parse_item_name(raw_name):
    raw_name = re.sub(r'\(\*[^)]*\)', '', raw_name)
    raw_name = re.sub(r'（\*[^）]*）', '', raw_name)
    match = re.search(r'[:：]\s*(.+)', clean_text(raw_name))
    if match:
        name = clean_text(match.group(1))
    else:
        name = clean_text(raw_name)
    return clean_chinese(name)


def parse_domain(raw_domain):
    """Returns (subcategory, 'PEF' or 'AMS')."""
    text = clean_text(raw_domain)
    text = re.sub(r'\(\s*\d+\s*\)\s*', '', text)
    text = clean_chinese(text)

    if '病理' in text:
        score_type = 'AMS'
        m = re.search(r'病理学?[：:]\s*(.+)', text)
        cat2 = clean_chinese(m.group(1)) if m else text.replace('病理学', '').replace('病理', '').strip('：: ')
    elif '发展' in text:
        score_type = 'PEF'
        m = re.search(r'发展[：:]\s*(.+)', text)
        cat2 = clean_chinese(m.group(1)) if m else text.replace('发展', '').strip('：: ')
    else:
        cat2 = clean_chinese(text)
        score_type = 'PEF'

    cat2 = cat2.replace('，', '、')
    cat2 = re.sub(r'\s+', '', cat2)
    if '人际关系' in cat2:
        cat2 = '人际关系、合作行为及对人的兴趣'
    return cat2 or '未分类', score_type


def parse_scoring(raw_scoring, score_type):
    """Returns (a_label, a_desc, b_label, b_desc, c_label, c_desc).
    a=highest(2), b=middle(1), c=lowest(0).
    """
    text = raw_scoring.strip()

    # Remove leading sub-domain prefix like "(2)手眼协调\n"
    if text.startswith('(') and re.match(r'\(\d\)', text):
        text = re.sub(r'^\([^)]*\)[^\n]*\n', '', text)

    if score_type == 'PEF':
        p_desc = e_desc = f_desc = ''
        pm = re.search(r'P[:：]\s*(.+?)(?=\nE[:：]|\Z)', text, re.DOTALL)
        if pm:
            p_desc = clean_text(pm.group(1))
        em = re.search(r'E[:：]\s*(.+?)(?=\nF[:：,i]|\Z)', text, re.DOTALL)
        if em:
            e_desc = clean_text(em.group(1))
        fm = re.search(r'F[i,：:]\s*(.+)', text, re.DOTALL)
        if fm:
            f_desc = clean_text(fm.group(1))
        return 'P', p_desc, 'E', e_desc, 'F', f_desc

    else:  # AMS
        a_desc = m_desc = s_desc = ''
        am = re.search(r'A[:：]\s*(.+?)(?=\nM[:：]|\Z)', text, re.DOTALL)
        if am:
            a_desc = clean_text(am.group(1))
        mm = re.search(r'M[:：]\s*(.+?)(?=\nS[:：]|\Z)', text, re.DOTALL)
        if mm:
            m_desc = clean_text(mm.group(1))
        sm = re.search(r'S[:：]\s*(.+)', text, re.DOTALL)
        if sm:
            s_desc = clean_text(sm.group(1))
        return 'A', a_desc, 'M', m_desc, 'S', s_desc


def main():
    print("Loading DOCX...")
    doc = Document(DOCX_FILE)

    # Extract valid test items from tables
    raw_items = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 5:
                continue
            if cells[0] == '任务项目' or len(set(cells)) == 1 or not cells[0].strip():
                continue
            domain = cells[3]
            scoring = cells[4]
            if ('发展' in domain or '病理' in domain
                    or re.search(r'[PAF][:：]', scoring)
                    or re.search(r'\(\s*\d\s*\)', domain)):
                raw_items.append(cells)

    print(f"  Raw items: {len(raw_items)}")

    # Extract supplementary items (补充项目) from tables 45-47
    supp_items = []
    current_cat2 = ''
    for table in doc.tables:
        first_cells = [cell.text.strip() for cell in table.rows[0].cells]
        if '补充项目' not in first_cells[0]:
            continue
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if r_idx == 0:  # header row
                continue
            if len(set(cells)) == 1:  # merged row = category header
                current_cat2 = clean_chinese(cells[0].strip())
                continue
            if not cells[0].strip():
                continue
            # Data row: [项目编号如"23(H)", 技能, 步骤]
            item_id = cells[0].strip()
            skill = clean_text(cells[1]) if len(cells) > 1 else ''
            procedure = clean_text(cells[2]) if len(cells) > 2 else ''
            supp_items.append({
                'item_id': item_id,
                'cat2': current_cat2,
                'skill': skill,
                'procedure': procedure,
            })

    print(f"  Supplementary items: {len(supp_items)}")

    # Build CSV rows - use original item numbers from DOCX
    csv_rows = []
    prev_name = ''
    prev_materials = ''
    prev_procedure = ''
    prev_base_num = ''
    current_seq = 0

    for cells in raw_items:
        item_name = parse_item_name(cells[0])
        materials = clean_text(cells[1])
        procedure = clean_text(cells[2])
        cat2, score_type = parse_domain(cells[3])

        # Extract original number from item name
        raw_text = re.sub(r'^[*·★°\s]+', '', clean_text(cells[0]))
        num_match = re.match(r'(\d+)\s*[A-Za-z]?\s*(?:-\s*\d+)?\s*[:：]', raw_text)
        if num_match:
            base_num = num_match.group(1)  # Just the main number (e.g. "8" from "8A")
        else:
            base_num = prev_base_num  # continuation row, keep previous

        # Continuation = same item with different scoring dimension
        is_cont = (item_name == prev_name
                   or bool(re.search(r'\(\s*2\s*\)', cells[3])))

        if not is_cont:
            current_seq = int(base_num) if base_num.isdigit() else current_seq + 1
            prev_name = item_name
            prev_materials = materials
            prev_procedure = procedure
            prev_base_num = base_num
        else:
            if not materials or materials == prev_materials:
                materials = prev_materials
            if not procedure or procedure == prev_procedure:
                procedure = prev_procedure

        # Extract original sub-number (e.g. "5A", "8B", "15A-1") for item name prefix
        raw_text_for_num = re.sub(r'^[*·★°\s]+', '', clean_text(cells[0]))
        sub_num_match = re.match(r'(\d+\s*[A-Za-z]?\s*(?:-\s*\d+)?)\s*[:：]', raw_text_for_num)
        if sub_num_match:
            sub_num = re.sub(r'\s+', '', sub_num_match.group(1))  # e.g. "5A", "15A-1"
        else:
            sub_num = base_num

        a_lbl, a_desc, b_lbl, b_desc, c_lbl, c_desc = parse_scoring(cells[4], score_type)

        csv_rows.append({
            "排序": str(current_seq),
            "一级分类": "病理学" if score_type == 'AMS' else "发展",
            "二级分类": cat2,
            "评估项目": f"{sub_num}：{item_name}",
            "操作描述": procedure,
            "所需材料": materials,
            "适用年龄": "",
            "是否必答": "是",
            "评分_a_分值": "2", "评分_a_标签": a_lbl, "评分_a_说明": a_desc,
            "评分_b_分值": "1", "评分_b_标签": b_lbl, "评分_b_说明": b_desc,
            "评分_c_分值": "0", "评分_c_标签": c_lbl, "评分_c_说明": c_desc,
            "评分_d_分值": "", "评分_d_标签": "", "评分_d_说明": "",
            "评分_e_分值": "", "评分_e_标签": "", "评分_e_说明": "",
        })

    # Write CSV
    # Append supplementary items - 排序从100开始自增
    supp_seq = 100
    for si in supp_items:
        item_id = si['item_id']  # e.g. "23(H)", "64(L)"
        skill = si['skill']      # e.g. "模仿", "模仿声音来获取奖励物"

        csv_rows.append({
            "排序": str(supp_seq),
            "一级分类": "补充项目",
            "二级分类": si['cat2'],
            "评估项目": f"{item_id}：{skill}",
            "操作描述": si['procedure'],
            "所需材料": "",
            "适用年龄": "",
            "是否必答": "否",
            "评分_a_分值": "1", "评分_a_标签": "H（通过）", "评分_a_说明": "",
            "评分_b_分值": "0", "评分_b_标签": "L（不通过）", "评分_b_说明": "",
            "评分_c_分值": "", "评分_c_标签": "", "评分_c_说明": "",
            "评分_d_分值": "", "评分_d_标签": "", "评分_d_说明": "",
            "评分_e_分值": "", "评分_e_标签": "", "评分_e_说明": "",
        })
        supp_seq += 1

    # Post-process: resolve "同X" material references
    # Build a lookup: sub_number -> materials (e.g. "1B" -> "一瓶肥皂泡液，吹泡泡的杆")
    mat_lookup = {}
    for r in csv_rows:
        if r['一级分类'] != '补充项目':
            # Extract sub-number from 评估项目 (e.g. "1B：吹泡泡" -> "1B")
            m = re.match(r'([^：]+)：', r['评估项目'])
            if m:
                key = m.group(1).strip()
                mat = r['所需材料'].strip()
                if mat and not mat.startswith('同'):
                    mat_lookup[key] = mat
                    # Also store without spaces for fuzzy matching
                    mat_lookup[re.sub(r'\s+', '', key)] = mat

    resolved_count = 0
    for r in csv_rows:
        mat = r['所需材料'].strip()
        if not mat.startswith('同'):
            continue

        # Parse reference: "同1 B", "同 3 8 A", "同 4 9 A , 但只使用颜色相同的积木"
        # Extract the referenced item number and optional suffix
        ref_match = re.match(r'同\s*(.+?)(?:\s*[,，]\s*(.+))?$', mat)
        if not ref_match:
            continue

        ref_num = re.sub(r'\s+', '', ref_match.group(1))  # e.g. "1B", "38A", "49A"
        suffix = ref_match.group(2) or ''  # e.g. "但只使用颜色相同的积木"

        # Look up the actual materials - try exact match first, then fuzzy
        actual_mat = mat_lookup.get(ref_num, '')
        if not actual_mat:
            # Try matching just the numeric part (e.g. "17A" -> look for any key starting with "17")
            base_digits = re.match(r'(\d+)', ref_num)
            if base_digits:
                digit_part = base_digits.group(1)
                # Find first key that starts with same digits and has materials
                for k, v in mat_lookup.items():
                    if re.match(rf'^{digit_part}[A-Za-z]?$', k) and v:
                        actual_mat = v
                        break
        if actual_mat:
            if suffix:
                r['所需材料'] = f"{actual_mat}（{suffix.strip()}）"
            else:
                r['所需材料'] = actual_mat
            resolved_count += 1

    print(f"  Resolved {resolved_count} '同X' material references")

    # Hard-fix remaining unresolved references (OCR number mismatches in DOCX)
    hardfix_map = {
        '17A': '一个大橡皮球',
        '38A': '一张汉字塑料板以及9张汉字卡（大、小、人、头、手、日、月、土、木）',
    }
    for r in csv_rows:
        mat = r['所需材料'].strip()
        if not mat.startswith('同'):
            continue
        ref_match = re.match(r'同\s*(.+?)(?:\s*[,，]\s*(.+))?$', mat)
        if not ref_match:
            continue
        ref_num = re.sub(r'\s+', '', ref_match.group(1))
        suffix = ref_match.group(2) or ''
        if ref_num in hardfix_map:
            if suffix:
                r['所需材料'] = f"{hardfix_map[ref_num]}（{suffix.strip()}）"
            else:
                r['所需材料'] = hardfix_map[ref_num]

    with open(OUTPUT, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=V2_HEADER, quoting=csv.QUOTE_ALL)
        writer.writeheader()
        writer.writerows(csv_rows)

    # Stats
    unique_seq = len(set(r['排序'] for r in csv_rows))
    pef = sum(1 for r in csv_rows if r['评分_a_标签'] == 'P')
    ams = sum(1 for r in csv_rows if r['评分_a_标签'] == 'A')
    has_a = sum(1 for r in csv_rows if r['评分_a_说明'])
    has_b = sum(1 for r in csv_rows if r['评分_b_说明'])
    has_c = sum(1 for r in csv_rows if r['评分_c_说明'])
    supp_count = sum(1 for r in csv_rows if r['一级分类'] == '补充项目')

    print(f"\nGenerated {OUTPUT}: {len(csv_rows)} rows")
    print(f"  Main items: {len(csv_rows) - supp_count} (P/E/F: {pef}, A/M/S: {ams})")
    print(f"  Supplementary items: {supp_count}")
    print(f"  评分说明 a/b/c: {has_a}/{has_b}/{has_c}")

    print("\nItems #59-60:")
    for r in csv_rows:
        if r['排序'] in ('59', '60'):
            print(f"  #{r['排序']} [{r['一级分类']}/{r['二级分类']}] {r['评估项目'][:30]}"
                  f" | a={r['评分_a_分值']}/{r['评分_a_标签']}"
                  f" b={r['评分_b_分值']}/{r['评分_b_标签']}"
                  f" c={r['评分_c_分值']}/{r['评分_c_标签']}")

    print("\nSample supplementary items:")
    for r in csv_rows:
        if r['一级分类'] == '补充项目':
            print(f"  #{r['排序']} [{r['二级分类']}] {r['评估项目']}"
                  f" | a={r['评分_a_分值']}/{r['评分_a_标签']}"
                  f" b={r['评分_b_分值']}/{r['评分_b_标签']}"
                  f" | {r['操作描述'][:40]}")
            if sum(1 for rr in csv_rows if rr['一级分类'] == '补充项目') > 10:
                break


if __name__ == "__main__":
    main()
